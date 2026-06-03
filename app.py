from flask import Flask, request, jsonify, send_file, send_from_directory, session
from flask_cors import CORS
import csv
import os
import io
import json
import requests
import uuid
import re
from werkzeug.utils import secure_filename
from datetime import datetime
from functools import wraps

app = Flask(__name__, static_folder='.', template_folder='.')
app.secret_key = os.environ.get('SECRET_KEY', 'ATM-platform-secret-key-2026-xK9mP2qL7vN4wR!')
CORS(app, supports_credentials=True)

app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024
ALLOWED_EXTENSIONS = {'csv', 'txt', 'pdf', 'docx', 'xlsx', 'xls', 'doc', 'rtf'}

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

uploaded_data = {}
USERS_FILE     = 'users.json'
FILES_META_FILE = 'files_meta.json'

PA_USERNAME = ''
PA_TOKEN    = ''

GIGACHAT_AUTH_KEY = ""


def load_users():
    if os.path.exists(USERS_FILE):
        with open(USERS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    default_users = {
        "admin": {"password": "admin", "role": "admin", "created": datetime.now().isoformat()}
    }
    save_users(default_users)
    return default_users


def save_users(users):
    with open(USERS_FILE, 'w', encoding='utf-8') as f:
        json.dump(users, f, ensure_ascii=False, indent=2)


def load_files_meta():
    if os.path.exists(FILES_META_FILE):
        with open(FILES_META_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}


def save_files_meta(meta):
    with open(FILES_META_FILE, 'w', encoding='utf-8') as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)


def restore_uploaded_data():
    meta = load_files_meta()
    for filename, info in meta.items():
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if os.path.exists(filepath):
            try:
                data = process_uploaded_file(filepath, filename)
                uploaded_data[filename] = {
                    'data': data,
                    'timestamp': info.get('timestamp', datetime.now().isoformat()),
                    'rows': len(data),
                    'uploader': info.get('uploader', 'unknown')
                }
            except Exception as e:
                print(f"[RESTORE] Не удалось восстановить {filename}: {e}")


def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user' not in session:
            return jsonify({"success": False, "message": "Требуется авторизация"}), 401
        return f(*args, **kwargs)
    return decorated


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user' not in session:
            return jsonify({"success": False, "message": "Требуется авторизация"}), 401
        users = load_users()
        if users.get(session['user'], {}).get('role') != 'admin':
            return jsonify({"success": False, "message": "Требуются права администратора"}), 403
        return f(*args, **kwargs)
    return decorated


@app.route('/api/login', methods=['POST'])
def login():
    body = request.json or {}
    username = body.get('username', '').strip()
    password = body.get('password', '').strip()
    if not username or not password:
        return jsonify({"success": False, "message": "Введите логин и пароль"}), 400
    users = load_users()
    if username in users and users[username]['password'] == password:
        session['user'] = username
        session.permanent = True
        return jsonify({"success": True, "username": username, "role": users[username].get('role', 'user')})
    return jsonify({"success": False, "message": "Неверный логин или пароль"}), 401


@app.route('/api/logout', methods=['POST'])
def logout():
    session.pop('user', None)
    return jsonify({"success": True})


@app.route('/api/me', methods=['GET'])
def me():
    if 'user' not in session:
        return jsonify({"logged_in": False})
    users = load_users()
    u = users.get(session['user'], {})
    return jsonify({"logged_in": True, "username": session['user'], "role": u.get('role', 'user')})


@app.route('/api/register', methods=['POST'])
@admin_required
def register():
    body = request.json or {}
    username = body.get('username', '').strip()
    password = body.get('password', '').strip()
    role = body.get('role', 'user')
    if not username or not password:
        return jsonify({"success": False, "message": "Введите логин и пароль"}), 400
    users = load_users()
    if username in users:
        return jsonify({"success": False, "message": "Пользователь уже существует"}), 400
    users[username] = {"password": password, "role": role, "created": datetime.now().isoformat()}
    save_users(users)
    return jsonify({"success": True, "message": f"Пользователь {username} создан"})


@app.route('/api/users', methods=['GET'])
@admin_required
def get_users():
    users = load_users()
    return jsonify({"success": True, "users": [
        {"username": u, "role": d.get('role', 'user'), "created": d.get('created', '')}
        for u, d in users.items()
    ]})


@app.route('/api/delete-user', methods=['POST'])
@admin_required
def delete_user():
    body = request.json or {}
    username = body.get('username', '')
    if username == session.get('user'):
        return jsonify({"success": False, "message": "Нельзя удалить себя"}), 400
    users = load_users()
    if username not in users:
        return jsonify({"success": False, "message": "Пользователь не найден"}), 404
    del users[username]
    save_users(users)
    return jsonify({"success": True})


@app.route('/api/change-password', methods=['POST'])
@login_required
def change_password():
    """
    Смена пароля.
    - Обычный пользователь: только свой пароль, обязан указать текущий.
    - Администратор: любой пользователь без текущего пароля.
    """
    body         = request.json or {}
    target_user  = body.get('username', '').strip()
    old_password = body.get('old_password', '').strip()
    new_password = body.get('new_password', '').strip()
    confirm      = body.get('confirm', '').strip()
    if not target_user or not new_password:
        return jsonify({"success": False, "message": "Укажите пользователя и новый пароль"}), 400
    if len(new_password) < 4:
        return jsonify({"success": False, "message": "Пароль должен быть не короче 4 символов"}), 400
    if new_password != confirm:
        return jsonify({"success": False, "message": "Пароли не совпадают"}), 400
    users   = load_users()
    me_name = session['user']
    me_role = users.get(me_name, {}).get('role', 'user')
    if target_user not in users:
        return jsonify({"success": False, "message": "Пользователь не найден"}), 404
    if me_role != 'admin':
        if target_user != me_name:
            return jsonify({"success": False, "message": "Нельзя менять чужой пароль"}), 403
        if not old_password:
            return jsonify({"success": False, "message": "Введите текущий пароль"}), 400
        if users[me_name]['password'] != old_password:
            return jsonify({"success": False, "message": "Неверный текущий пароль"}), 403
    else:
        if target_user == me_name and old_password:
            if users[me_name]['password'] != old_password:
                return jsonify({"success": False, "message": "Неверный текущий пароль"}), 403
    users[target_user]['password'] = new_password
    save_users(users)
    return jsonify({"success": True, "message": f"Пароль �у {target_user}� успешно изменён"})


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def detect_encoding(content_bytes):
    try:
        import chardet
        detected = chardet.detect(content_bytes)
        enc = (detected.get('encoding') or 'utf-8').lower()
    except Exception:
        enc = 'utf-8'
    if enc in ('windows-1251', 'cp1251', 'ansi', 'cyrillic'):
        return 'cp1251'
    if enc in ('koi8-r', 'koi8r'):
        return 'koi8-r'
    if enc in ('iso-8859-5', 'iso8859-5'):
        return 'iso-8859-5'
    return 'utf-8'


def decode_content(content_bytes):
    enc = detect_encoding(content_bytes)
    try:
        return content_bytes.decode(enc, errors='replace')
    except Exception:
        return content_bytes.decode('utf-8', errors='replace')


def find_8digit(text):
    m = re.search(r'\b(\d{8})\b', text)
    return m.group(1) if m else None


def strip_number(text, number):
    result = re.sub(r'\b' + re.escape(number) + r'\b', '', text)
    result = re.sub(r'^[\s,;|]+|[\s,;|]+$', '', result)
    return re.sub(r'\s{2,}', ' ', result).strip()


def make_record(number, address):
    city = ''
    if address:
        parts = address.split()
        city = parts[0] if parts else ''
    return {'number': number, 'address': address, 'city': city}


def deduplicate(records):
    seen = set()
    result = []
    for r in records:
        key = r.get('number') or r.get('text', '')
        if key and key not in seen:
            seen.add(key)
            result.append(r)
    return result


def _txt_method_tab(lines):
    data = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        for delim in ('\t', '|', ';', ','):
            if delim in line:
                parts = [p.strip() for p in line.split(delim, 1)]
                if len(parts) == 2:
                    num = find_8digit(parts[0]) or find_8digit(parts[1])
                    if num:
                        addr = parts[1] if find_8digit(parts[0]) else parts[0]
                        addr = strip_number(addr, num)
                        if addr:
                            data.append(make_record(num, addr))
                    break
    return data


def _txt_method_inline(lines):
    data = []
    for line in lines:
        line = line.strip()
        num = find_8digit(line)
        if num:
            addr = strip_number(line, num)
            if addr:
                data.append(make_record(num, addr))
    return data


def _txt_method_aggressive(text):
    data = []
    pattern = re.compile(
        r'(\d{8})\s+([\u0400-\u04FF\w\s\-,\(\)\.\'\"\/]{4,250}?)(?=\d{8}|\n\s*\n|$)',
        re.DOTALL
    )
    for m in pattern.finditer(text):
        num = m.group(1)
        addr = re.sub(r'\s+', ' ', m.group(2)).strip()
        addr = re.sub(r'^[,;\s]+|[,;\s]+$', '', addr)
        if addr and 4 < len(addr) < 300:
            data.append(make_record(num, addr))
    return data


def parse_txt_universal(content_bytes):
    content = decode_content(content_bytes).lstrip('\ufeff')
    lines = content.splitlines()
    skip_re = re.compile(r'номер|id|адрес|---|===|\*\*\*', re.IGNORECASE)
    data_lines = [l for l in lines if not skip_re.search(l)]
    results = []
    results.extend(_txt_method_tab(data_lines))
    results.extend(_txt_method_inline(data_lines))
    results.extend(_txt_method_aggressive(content))
    results = deduplicate(results)
    if not results:
        results = [{'line': i + 1, 'text': l.strip()} for i, l in enumerate(lines) if l.strip()]
    return results


def parse_docx_universal(file_path):
    import docx as _docx
    doc = _docx.Document(file_path)
    results = []
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if not cells:
                continue
            if len(cells) >= 2:
                num = find_8digit(cells[0])
                if num and cells[0].strip() == num:
                    addr = cells[1].strip()
                    if addr and addr.lower() not in ('адрес', 'адрес (организация)'):
                        results.append(make_record(num, addr))
                    continue
            row_text = ' '.join(cells)
            num = find_8digit(row_text)
            if num:
                if 'номер' in row_text.lower() or 'id' in row_text.lower():
                    continue
                addr = strip_number(row_text, num)
                if addr:
                    results.append(make_record(num, addr))
    skip_re = re.compile(r'номер|id|адрес|список|---|===', re.IGNORECASE)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or skip_re.search(text):
            continue
        num = find_8digit(text)
        if num:
            addr = strip_number(text, num)
            if addr:
                results.append(make_record(num, addr))
    full_parts = []
    for table in doc.tables:
        for row in table.rows:
            full_parts.append(' '.join(c.text for c in row.cells))
    for para in doc.paragraphs:
        full_parts.append(para.text)
    full_text = '\n'.join(full_parts)
    aggressive = _txt_method_aggressive(full_text)
    results.extend(aggressive)
    results = deduplicate(results)
    results.sort(key=lambda x: x.get('number', ''))
    return results


def _pdf_lines_from_page(page):
    text = page.extract_text() or ''
    return text.splitlines()


def parse_pdf_universal(file_path):
    results = []
    try:
        import pdfplumber
        full_text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                for line in _pdf_lines_from_page(page):
                    line = line.strip()
                    num = find_8digit(line)
                    if num:
                        addr = strip_number(line, num)
                        if addr:
                            results.append(make_record(num, addr))
                for table in (page.extract_tables() or []):
                    for row in table:
                        if not row:
                            continue
                        cells = [str(c).strip() for c in row if c]
                        if len(cells) >= 2:
                            num = find_8digit(cells[0])
                            if num and cells[0] == num:
                                results.append(make_record(num, cells[1]))
                                continue
                        row_text = ' '.join(cells)
                        num = find_8digit(row_text)
                        if num:
                            addr = strip_number(row_text, num)
                            if addr:
                                results.append(make_record(num, addr))
                text = page.extract_text() or ''
                full_text_parts.append(text)
        results.extend(_txt_method_aggressive('\n'.join(full_text_parts)))
    except Exception as e:
        print(f"[PDF] Ошибка: {e}", flush=True)
    return deduplicate(results)


def parse_excel_universal(file_path):
    results = []
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, data_only=True)
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if not cells:
                    continue
                if len(cells) >= 2:
                    num = find_8digit(cells[0])
                    if num and cells[0] == num:
                        results.append(make_record(num, cells[1]))
                        continue
                row_text = ' '.join(cells)
                num = find_8digit(row_text)
                if num:
                    addr = strip_number(row_text, num)
                    if addr:
                        results.append(make_record(num, addr))
    except Exception as e:
        print(f"[Excel] Ошибка: {e}", flush=True)
    return deduplicate(results)


def process_uploaded_file(file_path, filename):
    ext = filename.rsplit('.', 1)[-1].lower()
    if ext in ('csv', 'txt', 'rtf'):
        with open(file_path, 'rb') as f:
            return parse_txt_universal(f.read())
    elif ext == 'docx':
        return parse_docx_universal(file_path)
    elif ext == 'doc':
        try:
            import subprocess, tempfile, shutil
            tmp = tempfile.mkdtemp()
            subprocess.run(['soffice', '--headless', '--convert-to', 'docx',
                            '--outdir', tmp, file_path], timeout=60, capture_output=True)
            converted = os.path.join(tmp, os.path.splitext(os.path.basename(file_path))[0] + '.docx')
            if os.path.exists(converted):
                result = parse_docx_universal(converted)
                shutil.rmtree(tmp, ignore_errors=True)
                return result
        except Exception:
            pass
        with open(file_path, 'rb') as f:
            return parse_txt_universal(f.read())
    elif ext == 'pdf':
        return parse_pdf_universal(file_path)
    elif ext in ('xlsx', 'xls'):
        return parse_excel_universal(file_path)
    return []


def search_in_data(data, query):
    if not query or not data:
        return data[:500] if not query else []
    q = query.lower()
    results = []
    for item in data:
        item_str = json.dumps(item, ensure_ascii=False).lower()
        if q in item_str:
            results.append(item)
        if len(results) >= 500:
            break
    return results


def ask_gigachat(question, context_data=None, max_rows=500):
    if not question:
        return "Пожалуйста, задайте вопрос."

    try:
        import ssl
        ssl._create_default_https_context = ssl._create_unverified_context

        token_resp = requests.post(
            "https://ngw.devices.sberbank.ru:9443/api/v2/oauth",
            headers={
                'Content-Type': 'application/x-www-form-urlencoded',
                'Accept': 'application/json',
                'RqUID': str(uuid.uuid4()),
                'Authorization': f'Basic {GIGACHAT_AUTH_KEY}'
            },
            data="scope=GIGACHAT_API_PERS",
            verify=False, timeout=30
        )
        if token_resp.status_code != 200:
            return f"Ошибка авторизации GigaChat: {token_resp.status_code} — {token_resp.text[:200]}"
        access_token = token_resp.json().get('access_token')

        chat_url = "https://gigachat.devices.sberbank.ru/api/v1/chat/completions"
        chat_headers = {
            'Content-Type': 'application/json',
            'Accept': 'application/json',
            'Authorization': f'Bearer {access_token}',
            'RqUID': str(uuid.uuid4())
        }
        system_prompt = (
            "Ты — ИИ-ассистент информационной платформы. "
            "Отвечай на вопросы по содержимому файлов, помогай находить банкоматы, "
            "анализируй данные. Будь вежливым и полезным."
        )

        if not context_data or len(context_data) <= 200:
            user_message = question
            if context_data:
                preview = json.dumps(context_data[:200], ensure_ascii=False, indent=2)[:8000]
                user_message = (
                    f"Контекст из загруженного файла ({len(context_data)} записей):\n{preview}\n\n"
                    f"Вопрос пользователя: {question}"
                )
            payload = {
                "model": "GigaChat",
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user",   "content": user_message}
                ],
                "temperature": 0.7,
                "max_tokens": 3000
            }
            resp = requests.post(chat_url, headers=chat_headers,
                                 json=payload, verify=False, timeout=90)
            if resp.status_code == 200:
                return (resp.json().get('choices', [{}])[0]
                        .get('message', {}).get('content', 'Нет ответа'))
            return f"Ошибка GigaChat {resp.status_code}: {resp.text[:300]}"

        chunk_size = 200
        chunks = [context_data[i:i + chunk_size]
                  for i in range(0, min(len(context_data), max_rows), chunk_size)]

        partial_answers = []
        for idx, chunk in enumerate(chunks):
            preview = json.dumps(chunk, ensure_ascii=False, indent=2)[:6000]
            chunk_message = (
                f"Это часть {idx + 1} из {len(chunks)} большого файла "
                f"(всего {len(context_data)} записей, строки {idx*chunk_size+1}–{idx*chunk_size+len(chunk)}).\n"
                f"Данные:\n{preview}\n\n"
                f"Вопрос пользователя: {question}\n\n"
                f"Дай краткий промежуточный ответ по этой части. "
                f"Если релевантных данных нет — напиши 'нет данных в этой части'."
            )
            payload = {
                "model": "GigaChat",
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user",   "content": chunk_message}
                ],
                "temperature": 0.5,
                "max_tokens": 1500
            }
            resp = requests.post(chat_url, headers=chat_headers,
                                 json=payload, verify=False, timeout=90)
            if resp.status_code == 200:
                part = (resp.json().get('choices', [{}])[0]
                        .get('message', {}).get('content', ''))
                if part and 'нет данных' not in part.lower():
                    partial_answers.append(f"[Часть {idx + 1}]: {part}")

        if not partial_answers:
            return "По вашему запросу ничего не найдено в загруженном файле."

        synthesis_message = (
            f"Я проанализировал файл из {len(context_data)} строк по частям. "
            f"Промежуточные результаты:\n\n"
            + "\n\n".join(partial_answers)
            + f"\n\nВопрос: {question}\n\n"
            f"На основе всех частей дай итоговый структурированный ответ."
        )
        payload = {
            "model": "GigaChat",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": synthesis_message}
            ],
            "temperature": 0.7,
            "max_tokens": 3000
        }
        resp = requests.post(chat_url, headers=chat_headers,
                             json=payload, verify=False, timeout=90)
        if resp.status_code == 200:
            return (resp.json().get('choices', [{}])[0]
                    .get('message', {}).get('content', 'Нет ответа'))
        return f"Ошибка финального синтеза {resp.status_code}: {resp.text[:300]}"

    except Exception as e:
        return f"Ошибка при обращении к GigaChat: {e}"


@app.route('/')
def index():
    return send_from_directory('.', 'index.html')


@app.route('/style.css')
def serve_css():
    return send_from_directory('.', 'style.css')


@app.route('/api/upload', methods=['POST'])
@login_required
def upload_file():
    if 'file' not in request.files:
        return jsonify({"success": False, "message": "Файл не выбран"}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({"success": False, "message": "Файл не выбран"}), 400
    if not allowed_file(file.filename):
        return jsonify({"success": False,
                        "message": f"Неподдерживаемый формат. Разрешены: {', '.join(ALLOWED_EXTENSIONS)}"}), 400
    try:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        data = process_uploaded_file(filepath, filename)
        ts = datetime.now().isoformat()
        uploaded_data[filename] = {
            'data': data, 'timestamp': ts,
            'rows': len(data), 'uploader': session.get('user', 'unknown')
        }
        meta = load_files_meta()
        meta[filename] = {'timestamp': ts, 'rows': len(data), 'uploader': session.get('user', 'unknown')}
        save_files_meta(meta)
        return jsonify({
            "success": True, "filename": filename, "rows": len(data),
            "preview": data[:20],
            "message": f"Файл {filename} успешно загружен. Обработано записей: {len(data)}",
            "can_import": len(data) > 0
        })
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"success": False, "message": f"Ошибка: {e}"}), 500


@app.route('/api/delete-file', methods=['POST'])
@login_required
def delete_file():
    body = request.json or {}
    filename = body.get('filename', '').strip()
    if not filename:
        return jsonify({"success": False, "message": "Укажите файл"}), 400
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(filepath):
        os.remove(filepath)
    uploaded_data.pop(filename, None)
    meta = load_files_meta()
    meta.pop(filename, None)
    save_files_meta(meta)
    return jsonify({"success": True, "message": f"Файл {filename} удалён"})


@app.route('/api/get-uploaded-files', methods=['GET'])
@login_required
def get_uploaded_files():
    files_list = [
        {"name": name, "rows": d['rows'], "date": d['timestamp'], "uploader": d.get('uploader', '—')}
        for name, d in uploaded_data.items()
    ]
    return jsonify({"success": True, "files": files_list})


@app.route('/api/search-in-file', methods=['POST'])
@login_required
def search_in_uploaded_file():
    body = request.json or {}
    filename = body.get('filename', '')
    query    = body.get('query', '')
    if not filename:
        return jsonify({"success": False, "message": "Укажите файл"}), 400
    if filename not in uploaded_data:
        return jsonify({"success": False, "message": "Файл не найден на сервере"}), 404
    results = search_in_data(uploaded_data[filename]['data'], query)
    return jsonify({
        "success": True, "data": results, "count": len(results),
        "filename": filename, "total": uploaded_data[filename]['rows']
    })


@app.route('/api/upload-yandex', methods=['POST'])
@login_required
def upload_from_yandex():
    body = request.json or {}
    public_url = body.get('url', '').strip()
    if not public_url:
        return jsonify({"success": False, "message": "Введите публичную ссылку"}), 400
    try:
        api_url = f"https://cloud-api.yandex.net/v1/disk/public/resources?public_key={public_url}"
        resp = requests.get(api_url, timeout=30)
        if resp.status_code != 200:
            return jsonify({"success": False,
                            "message": f"Не удалось получить информацию о файле (код {resp.status_code}). "
                                       "Проверьте ссылку и убедитесь, что файл открыт для всех."}), 400
        file_info = resp.json()
        if file_info.get('type') == 'dir':
            return jsonify({"success": False,
                            "message": "Это папка, а не файл. Укажите прямую ссылку на файл."}), 400

        dl_resp = requests.get(
            f"https://cloud-api.yandex.net/v1/disk/public/resources/download?public_key={public_url}",
            timeout=30
        )
        if dl_resp.status_code != 200:
            return jsonify({"success": False, "message": "Не удалось получить ссылку для скачивания"}), 400
        download_url = dl_resp.json().get('href')
        if not download_url:
            return jsonify({"success": False, "message": "Пустая ссылка на скачивание"}), 400

        file_resp = requests.get(download_url, timeout=120)
        filename      = file_info.get('name', 'yandex_file')
        content_bytes = file_resp.content

        ext      = filename.rsplit('.', 1)[-1].lower()
        tmp_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        with open(tmp_path, 'wb') as f:
            f.write(content_bytes)

        if ext in ('csv', 'txt', 'rtf'):
            data = parse_txt_universal(content_bytes)
        elif ext in ('docx', 'pdf', 'xlsx', 'xls', 'doc'):
            data = process_uploaded_file(tmp_path, filename)
        else:
            data = [{'text': l.strip()} for l in
                    content_bytes.decode('utf-8', errors='replace').splitlines() if l.strip()]

        ts = datetime.now().isoformat()
        uploaded_data[filename] = {
            'data': data, 'timestamp': ts,
            'rows': len(data), 'uploader': session.get('user', 'unknown'), 'source': 'yandex'
        }
        meta = load_files_meta()
        meta[filename] = {'timestamp': ts, 'rows': len(data),
                          'uploader': session.get('user', ''), 'source': 'yandex'}
        save_files_meta(meta)

        return jsonify({
            "success": True, "filename": filename, "rows": len(data),
            "preview": data[:20],
            "message": f"Файл {filename} загружен с Яндекс.Диска. Записей: {len(data)}",
            "can_import": len(data) > 0
        })
    except Exception as e:
        return jsonify({"success": False, "message": f"Ошибка: {e}"}), 500


@app.route('/api/gigachat', methods=['POST'])
@login_required
def gigachat_endpoint():
    body     = request.json or {}
    question = body.get('question', '')
    filename = body.get('filename', '')
    max_rows = int(body.get('max_rows', 500))

    if not question:
        return jsonify({"success": False, "message": "Введите вопрос"}), 400

    context_data = None
    if filename and filename in uploaded_data:
        context_data = uploaded_data[filename]['data']
    elif uploaded_data:
        context_data = list(uploaded_data.values())[-1]['data']

    answer   = ask_gigachat(question, context_data, max_rows=max_rows)
    rows_used = len(context_data) if context_data else 0
    return jsonify({"success": True, "answer": answer, "rows_analyzed": rows_used})


@app.route('/api/export', methods=['POST'])
@login_required
def export_results():
    body     = request.json or {}
    data     = body.get('data', [])
    filename = body.get('filename', 'export')
    if not data:
        return jsonify({"success": False, "message": "Нет данных"}), 400
    output = io.StringIO()
    if data and isinstance(data[0], dict):
        writer = csv.DictWriter(output, fieldnames=data[0].keys())
        writer.writeheader()
        writer.writerows(data)
    else:
        writer = csv.writer(output)
        for item in data:
            writer.writerow([str(item)])
    output.seek(0)
    ts      = datetime.now().strftime('%Y%m%d_%H%M%S')
    dl_name = f"{filename.rsplit('.', 1)[0]}_{ts}.csv"
    return send_file(
        io.BytesIO(output.getvalue().encode('utf-8-sig')),
        mimetype='text/csv', as_attachment=True, download_name=dl_name
    )


@app.route('/api/extend-site', methods=['POST'])
@admin_required
def extend_site():
    """Перезагружает и продлевает веб-приложение на PythonAnywhere."""
    try:
        cpu_resp = requests.get(
            f"https://www.pythonanywhere.com/api/v0/user/{PA_USERNAME}/cpu/",
            headers={'Authorization': f'Token {PA_TOKEN}'},
            timeout=15
        )
        if cpu_resp.status_code != 200:
            return jsonify({
                "success": False,
                "message": f"Ошибка авторизации PythonAnywhere (код {cpu_resp.status_code}). Проверьте токен."
            }), 400

        reload_url = (
            f"https://www.pythonanywhere.com/api/v0/user/{PA_USERNAME}"
            f"/webapps/{PA_USERNAME}.pythonanywhere.com/reload/"
        )
        resp = requests.post(reload_url, headers={'Authorization': f'Token {PA_TOKEN}'}, timeout=30)
        if resp.status_code == 200:
            return jsonify({
                "success": True,
                "message": f"✅ Сайт {PA_USERNAME}.pythonanywhere.com перезагружен и продлён!"
            })
        return jsonify({
            "success": False,
            "message": f"Ошибка перезагрузки: код {resp.status_code} — {resp.text[:200]}"
        })
    except Exception as e:
        return jsonify({"success": False, "message": f"Ошибка: {e}"}), 500


restore_uploaded_data()
application = app

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
